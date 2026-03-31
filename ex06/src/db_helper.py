"""ex06 — DB 연결 및 벡터스토어 헬퍼."""

import os
from pathlib import Path

DB_ERROR_MSG = (
    "PostgreSQL에 연결할 수 없습니다. "
    "docker-compose up -d 로 DB를 시작한 후 다시 시도하십시오."
)


def run_query(sql, params=()):
    """PostgreSQL 쿼리를 실행하고 결과를 반환한다."""
    try:
        import psycopg2
        import psycopg2.extras

        db_url = os.getenv("DATABASE_URL", "")
        if not db_url:
            host = os.getenv("POSTGRES_HOST", "localhost")
            port = os.getenv("POSTGRES_PORT", "5432")
            db = os.getenv("POSTGRES_DB", "rag_db")
            user = os.getenv("POSTGRES_USER", "rag_user")
            password = os.getenv("POSTGRES_PASSWORD", "")
            db_url = f"postgresql://{user}:{password}@{host}:{port}/{db}"

        conn = psycopg2.connect(db_url, connect_timeout=3)
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(sql, params)
        rows = [dict(row) for row in cur.fetchall()]
        cur.close()
        conn.close()
        return rows
    except Exception:
        return []


def build_vectorstore():
    """ChromaDB VectorStore를 생성한다. 없으면 data/docs/에서 자동 구축한다."""
    try:
        import chromadb
        from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
    except ImportError:
        print("[경고] chromadb 또는 sentence-transformers 미설치. 키워드 검색으로 동작합니다.")
        return None

    chroma_path = os.getenv("CHROMA_PERSIST_DIR", "./data/chroma_db")
    collection_name = os.getenv("CHROMA_COLLECTION_NAME", "metacoding_documents")
    ef = SentenceTransformerEmbeddingFunction(model_name="jhgan/ko-sroberta-multitask")

    client = chromadb.PersistentClient(path=chroma_path)

    try:
        collection = client.get_collection(collection_name, embedding_function=ef)
        if collection.count() > 0:
            print(f"[INFO] 기존 ChromaDB 로드: {chroma_path} ({collection.count()}건)")
            return collection
    except Exception:
        pass

    print("[INFO] ChromaDB가 없습니다. data/docs/ 원본 문서에서 자동 구축합니다.")
    docs = parse_and_chunk_docs()
    if not docs:
        print("[경고] data/docs/에 문서가 없습니다. 키워드 검색으로 동작합니다.")
        return None

    collection = client.get_or_create_collection(collection_name, embedding_function=ef)
    for i, doc in enumerate(docs):
        collection.add(
            ids=[f"doc_{i}"],
            documents=[doc["content"]],
            metadatas=[{"source": doc["source"], "page": doc.get("page", 1)}],
        )
    print(f"[INFO] ChromaDB 자동 구축 완료: {len(docs)}건 → {chroma_path}")
    return collection


def parse_and_chunk_docs(chunk_size=500, overlap=100):
    """data/docs/의 원본 PDF/DOCX/XLSX를 파싱→청킹하여 딕셔너리 리스트로 반환한다."""
    import pypdf
    from docx import Document as DocxDocument
    import openpyxl

    docs_dir = Path(__file__).parent.parent / "data" / "docs"
    if not docs_dir.exists():
        return []

    documents = []
    for file_path in sorted(docs_dir.rglob("*")):
        suffix = file_path.suffix.lower()
        source = file_path.stem
        texts = []

        if suffix == ".pdf":
            try:
                with open(file_path, "rb") as f:
                    reader = pypdf.PdfReader(f)
                    for page_num, page in enumerate(reader.pages, start=1):
                        text = (page.extract_text() or "").strip()
                        if text:
                            md_text = f"## 페이지 {page_num}\n\n{text}"
                            texts.append((md_text, page_num))
                print(f"  [INFO] PDF 파싱: {file_path.name} ({len(reader.pages)}페이지)")
            except Exception:
                continue
        elif suffix == ".docx":
            try:
                doc = DocxDocument(str(file_path))
                text_parts = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    style_name = para.style.name
                    if style_name == "Title":
                        text_parts.append(f"# {text}")
                    elif style_name.startswith("Heading"):
                        level_str = style_name.replace("Heading", "").strip()
                        try:
                            level = int(level_str)
                        except ValueError:
                            level = 2
                        text_parts.append(f"{'#' * level} {text}")
                    elif style_name == "List Bullet":
                        text_parts.append(f"- {text}")
                    else:
                        text_parts.append(text)
                for table in doc.tables:
                    for i, row in enumerate(table.rows):
                        row_data = [cell.text.strip() for cell in row.cells]
                        text_parts.append("| " + " | ".join(row_data) + " |")
                        if i == 0:
                            text_parts.append("| " + " | ".join(["---"] * len(row_data)) + " |")
                full_text = "\n".join(text_parts)
                if full_text:
                    texts.append((full_text, 1))
                print(f"  [INFO] DOCX 파싱: {file_path.name}")
            except Exception:
                continue
        elif suffix == ".xlsx":
            try:
                wb = openpyxl.load_workbook(str(file_path), data_only=True)
                for idx, name in enumerate(wb.sheetnames, start=1):
                    ws = wb[name]
                    rows = []
                    for row in ws.iter_rows():
                        cell_values = [
                            str(c.value).strip()
                            for c in row
                            if c.value is not None and str(c.value).strip()
                        ]
                        if cell_values:
                            rows.append(cell_values)
                    if rows:
                        max_col = max(len(r) for r in rows)
                        md_lines = [f"[시트: {name}]"]
                        header = rows[0] + [""] * (max_col - len(rows[0]))
                        md_lines.append("| " + " | ".join(header) + " |")
                        md_lines.append("| " + " | ".join(["---"] * max_col) + " |")
                        for row_data in rows[1:]:
                            row_data = row_data + [""] * (max_col - len(row_data))
                            md_lines.append("| " + " | ".join(row_data) + " |")
                        texts.append(("\n".join(md_lines), idx))
                print(f"  [INFO] XLSX 파싱: {file_path.name} ({len(wb.sheetnames)}시트)")
            except Exception:
                continue

        for text, page_num in texts:
            step = chunk_size - overlap
            start = 0
            while start < len(text):
                chunk = text[start:start + chunk_size].strip()
                if chunk:
                    documents.append({"content": chunk, "source": source, "page": page_num})
                start += step

    return documents


# 모듈 레벨 벡터스토어 캐시 (한 번만 구축)
_VECTORSTORE_CACHE = None


def get_vectorstore():
    """벡터스토어 싱글톤을 반환한다."""
    global _VECTORSTORE_CACHE
    if _VECTORSTORE_CACHE is None:
        _VECTORSTORE_CACHE = build_vectorstore()
    return _VECTORSTORE_CACHE
