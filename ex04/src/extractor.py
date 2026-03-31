"""
PDF, DOCX, XLSX 파일에서 텍스트를 추출하는 모듈.

Python 파싱 라이브러리(pypdf, python-docx, openpyxl)를 사용하여
다양한 형식의 문서에서 순수 텍스트를 추출하고, 추출 결과를
표준화된 딕셔너리 형식으로 반환합니다.
"""

import sys
from pathlib import Path

import openpyxl
import pypdf
from docx import Document


# =====================================================================
# === INPUT ===
# file_path: 추출할 문서 파일의 경로 (str 또는 Path)
# =====================================================================


def extract_from_pdf(file_path: str | Path) -> dict:
    """PDF 파일에서 텍스트를 페이지별로 추출합니다.

    pypdf를 사용하여 각 페이지의 텍스트를 순서대로 추출합니다.
    이미지 기반 PDF나 복잡한 레이아웃(다단, 표 안 이미지)은
    텍스트 손실이 발생할 수 있습니다.

    Args:
        file_path: 추출할 PDF 파일 경로

    Returns:
        {
            "source_path": 파일 절대 경로 (str),
            "file_name": 파일명 (str),
            "file_type": "pdf",
            "pages": [{"page": 페이지 번호, "text": 추출 텍스트}, ...],
            "full_text": 전체 텍스트 연결 문자열 (str)
        }

    Raises:
        FileNotFoundError: 파일이 존재하지 않을 경우
        RuntimeError: PDF 파싱 중 오류가 발생한 경우
    """
    file_path = Path(file_path)
    if not file_path.exists():
        print(f"파일을 찾을 수 없습니다: {file_path}")
        print("파일 경로를 확인하십시오.")
        sys.exit(1)

    # === PROCESS ===
    pages_data = []
    try:
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    md_text = f"## 페이지 {page_num}\n\n{page_text}"
                    pages_data.append({"page": page_num, "text": md_text})
                else:
                    pages_data.append({"page": page_num, "text": ""})
    except Exception as e:
        raise RuntimeError(
            f"PDF 파싱 중 오류가 발생했습니다: {file_path.name}\n원인: {e}"
        ) from e

    full_text = "\n\n".join(p["text"] for p in pages_data if p["text"])

    # === OUTPUT ===
    return {
        "source_path": str(file_path.resolve()),
        "file_name": file_path.name,
        "file_type": "pdf",
        "pages": pages_data,
        "full_text": full_text,
    }


def extract_from_docx(file_path: str | Path) -> dict:
    """DOCX 파일에서 단락과 표의 텍스트를 추출합니다.

    python-docx를 사용하여 워드 문서의 단락(Paragraph)과
    표(Table) 내용을 추출합니다. 제목(Heading) 스타일은
    마크다운 헤더 형식(#, ##)으로 변환됩니다.

    Args:
        file_path: 추출할 DOCX 파일 경로

    Returns:
        {
            "source_path": 파일 절대 경로 (str),
            "file_name": 파일명 (str),
            "file_type": "docx",
            "pages": [{"page": 1, "text": 전체 텍스트}],
            "full_text": 전체 텍스트 (str)
        }

    Raises:
        FileNotFoundError: 파일이 존재하지 않을 경우
        RuntimeError: DOCX 파싱 중 오류가 발생한 경우
    """
    file_path = Path(file_path)
    if not file_path.exists():
        print(f"파일을 찾을 수 없습니다: {file_path}")
        print("파일 경로를 확인하십시오.")
        sys.exit(1)

    # === PROCESS ===
    text_parts = []
    try:
        doc = Document(str(file_path))

        # 단락(Paragraph) 추출: 제목 스타일은 마크다운 헤더로 변환
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name
            if style_name == "Title":
                text_parts.append(f"# {text}")
            elif style_name.startswith("Heading"):
                level_str = style_name.replace("Heading", "").strip()
                try:
                    level = int(level_str)
                except ValueError:
                    level = 2
                text_parts.append(f"{'#' * level} {text}")
            elif style_name == "List Bullet":
                text_parts.append(f"- {text}")
            else:
                text_parts.append(text)

        # 표(Table) 추출: 행×열 구조를 마크다운 표로 변환
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                text_parts.append("| " + " | ".join(row_data) + " |")
                if i == 0:
                    text_parts.append("| " + " | ".join(["---"] * len(row_data)) + " |")

    except Exception as e:
        raise RuntimeError(
            f"DOCX 파싱 중 오류가 발생했습니다: {file_path.name}\n원인: {e}"
        ) from e

    full_text = "\n".join(text_parts)

    # === OUTPUT ===
    return {
        "source_path": str(file_path.resolve()),
        "file_name": file_path.name,
        "file_type": "docx",
        "pages": [{"page": 1, "text": full_text}],
        "full_text": full_text,
    }


def extract_from_xlsx(file_path: str | Path) -> dict:
    """XLSX 파일에서 각 시트의 셀 텍스트를 추출합니다.

    openpyxl을 사용하여 엑셀 파일의 모든 시트에서
    비어 있지 않은 셀 값을 읽어 행 단위로 연결합니다.
    병합 셀의 경우 openpyxl이 자동으로 최상단 셀 값을 반환합니다.

    Args:
        file_path: 추출할 XLSX 파일 경로

    Returns:
        {
            "source_path": 파일 절대 경로 (str),
            "file_name": 파일명 (str),
            "file_type": "xlsx",
            "pages": [{"page": 시트 순번, "text": 시트 텍스트}, ...],
            "full_text": 전체 텍스트 연결 문자열 (str)
        }

    Raises:
        FileNotFoundError: 파일이 존재하지 않을 경우
        RuntimeError: XLSX 파싱 중 오류가 발생한 경우
    """
    file_path = Path(file_path)
    if not file_path.exists():
        print(f"파일을 찾을 수 없습니다: {file_path}")
        print("파일 경로를 확인하십시오.")
        sys.exit(1)

    # === PROCESS ===
    pages_data = []
    try:
        wb = openpyxl.load_workbook(str(file_path), data_only=True)
        for sheet_idx, sheet_name in enumerate(wb.sheetnames, start=1):
            ws = wb[sheet_name]

            # 모든 행을 리스트로 수집
            rows = []
            for row in ws.iter_rows():
                cell_values = [
                    str(cell.value).strip()
                    for cell in row
                    if cell.value is not None and str(cell.value).strip()
                ]
                if cell_values:
                    rows.append(cell_values)

            if not rows:
                continue

            # 마크다운 표 형식으로 변환
            max_col = max(len(r) for r in rows)
            md_lines = [f"[시트: {sheet_name}]"]

            # 첫 번째 행 = 헤더
            header = rows[0] + [""] * (max_col - len(rows[0]))
            md_lines.append("| " + " | ".join(header) + " |")
            md_lines.append("| " + " | ".join(["---"] * max_col) + " |")

            # 나머지 행 = 데이터
            for row_data in rows[1:]:
                row_data = row_data + [""] * (max_col - len(row_data))
                md_lines.append("| " + " | ".join(row_data) + " |")

            sheet_text = "\n".join(md_lines)
            pages_data.append({"page": sheet_idx, "text": sheet_text})
    except Exception as e:
        raise RuntimeError(
            f"XLSX 파싱 중 오류가 발생했습니다: {file_path.name}\n원인: {e}"
        ) from e

    full_text = "\n\n".join(p["text"] for p in pages_data if p["text"])

    # === OUTPUT ===
    return {
        "source_path": str(file_path.resolve()),
        "file_name": file_path.name,
        "file_type": "xlsx",
        "pages": pages_data,
        "full_text": full_text,
    }


def extract_text(file_path: str | Path) -> dict:
    """파일 형식을 자동 감지하여 텍스트를 추출하는 통합 함수.

    파일 확장자를 기준으로 PDF, DOCX, XLSX 파싱 함수 중
    적절한 것을 선택하여 호출합니다.

    Args:
        file_path: 추출할 문서 파일 경로 (PDF/DOCX/XLSX)

    Returns:
        extract_from_pdf / extract_from_docx / extract_from_xlsx 중
        해당 형식의 반환값과 동일한 구조의 딕셔너리

    Raises:
        ValueError: 지원하지 않는 파일 형식인 경우
        FileNotFoundError: 파일이 존재하지 않을 경우
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    # === PROCESS ===
    extractor_map = {
        ".pdf": extract_from_pdf,
        ".docx": extract_from_docx,
        ".xlsx": extract_from_xlsx,
    }

    if suffix not in extractor_map:
        raise ValueError(
            f"지원하지 않는 파일 형식입니다: '{suffix}'\n"
            f"지원 형식: {', '.join(extractor_map.keys())}"
        )

    extractor_fn = extractor_map[suffix]

    # === OUTPUT ===
    return extractor_fn(file_path)


def extract_all_from_directory(docs_dir: str | Path) -> list[dict]:
    """디렉토리 내 모든 지원 문서를 재귀적으로 추출합니다.

    docs_dir 하위의 모든 PDF, DOCX, XLSX 파일을 찾아
    extract_text()를 호출하고 결과를 리스트로 반환합니다.
    파싱에 실패한 파일은 경고 메시지를 출력하고 건너뜁니다.

    Args:
        docs_dir: 문서가 저장된 최상위 디렉토리 경로

    Returns:
        각 문서의 추출 결과 딕셔너리 리스트
    """
    docs_dir = Path(docs_dir)
    if not docs_dir.exists():
        print(f"문서 디렉토리를 찾을 수 없습니다: {docs_dir}")
        print("data/docs/ 폴더에 문서 파일을 넣고 다시 실행하십시오.")
        sys.exit(1)

    # === PROCESS ===
    supported_extensions = {".pdf", ".docx", ".xlsx"}
    results = []

    all_files = sorted(
        [f for f in docs_dir.rglob("*") if f.suffix.lower() in supported_extensions]
    )

    if not all_files:
        print(f"지원 형식의 문서가 없습니다: {docs_dir}")
        return results

    print(f"총 {len(all_files)}개 문서를 발견했습니다.")

    for file_path in all_files:
        print(f"  추출 중: {file_path.name} ...", end=" ", flush=True)
        try:
            result = extract_text(file_path)
            results.append(result)
            text_length = len(result["full_text"])
            print(f"완료 ({text_length}자)")
        except Exception as e:
            print(f"실패 — {e}")

    # === OUTPUT ===
    return results
