"""ex10 문서 검색 도구."""

import os
import logging
from pathlib import Path

from langchain_core.tools import tool

logger = logging.getLogger(__name__)


def _build_vectorstore():
    """ChromaDB Collection을 생성한다. 없으면 data/docs/에서 자동 구축한다."""
    try:
        import chromadb
        from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
    except ImportError:
        logger.warning("chromadb 또는 sentence-transformers 미설치. 키워드 검색으로 동작합니다.")
        return None

    chroma_dir = os.getenv("CHROMA_PERSIST_DIR", "./data/chroma_db")
    collection_name = os.getenv("CHROMA_COLLECTION_NAME", "metacoding_documents")
    ef = SentenceTransformerEmbeddingFunction(model_name="jhgan/ko-sroberta-multitask")

    client = chromadb.PersistentClient(path=chroma_dir)

    # 기존 ChromaDB 데이터 확인
    try:
        collection = client.get_collection(collection_name, embedding_function=ef)
        if collection.count() > 0:
            logger.info("기존 ChromaDB 로드: %s (%d건)", chroma_dir, collection.count())
            return collection
    except Exception:
        pass

    # ChromaDB 없음 → data/docs/ 원본 문서에서 자동 구축
    logger.info("ChromaDB가 없습니다. data/docs/ 원본 문서에서 자동 구축합니다.")
    docs = _parse_and_chunk_docs()
    if not docs:
        logger.warning("data/docs/에 문서가 없습니다. 키워드 검색으로 동작합니다.")
        return None

    collection = client.get_or_create_collection(collection_name, embedding_function=ef)
    for i, doc in enumerate(docs):
        collection.add(
            ids=[f"doc_{i}"],
            documents=[doc["content"]],
            metadatas=[{"source": doc["source"], "page": doc.get("page", 1)}],
        )
    logger.info("ChromaDB 자동 구축 완료: %d건 → %s", len(docs), chroma_dir)
    return collection


def _parse_and_chunk_docs(chunk_size=500, overlap=100):
    """data/docs/의 원본 PDF/DOCX/XLSX를 파싱하여 딕셔너리 리스트로 반환한다."""
    import pypdf
    from docx import Document as DocxDocument
    import openpyxl

    docs_dir = Path(__file__).resolve().parent.parent.parent / "data" / "docs"
    if not docs_dir.exists():
        return []

    documents = []
    for file_path in sorted(docs_dir.rglob("*")):
        suffix = file_path.suffix.lower()
        source = file_path.stem
        texts = []

        if suffix == ".pdf":
            try:
                with open(file_path, "rb") as f:
                    reader = pypdf.PdfReader(f)
                    for page_num, page in enumerate(reader.pages, start=1):
                        text = (page.extract_text() or "").strip()
                        if text:
                            md_text = f"## 페이지 {page_num}\n\n{text}"
                            texts.append((md_text, page_num))
            except Exception:
                continue
        elif suffix == ".docx":
            try:
                doc = DocxDocument(str(file_path))
                text_parts = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    style_name = para.style.name
                    if style_name == "Title":
                        text_parts.append(f"# {text}")
                    elif style_name.startswith("Heading"):
                        level_str = style_name.replace("Heading", "").strip()
                        try:
                            level = int(level_str)
                        except ValueError:
                            level = 2
                        text_parts.append(f"{'#' * level} {text}")
                    elif style_name == "List Bullet":
                        text_parts.append(f"- {text}")
                    else:
                        text_parts.append(text)
                for table in doc.tables:
                    for i, row in enumerate(table.rows):
                        row_data = [cell.text.strip() for cell in row.cells]
                        text_parts.append("| " + " | ".join(row_data) + " |")
                        if i == 0:
                            text_parts.append("| " + " | ".join(["---"] * len(row_data)) + " |")
                full_text = "\n".join(text_parts)
                if full_text:
                    texts.append((full_text, 1))
            except Exception:
                continue
        elif suffix == ".xlsx":
            try:
                wb = openpyxl.load_workbook(str(file_path), data_only=True)
                for idx, name in enumerate(wb.sheetnames, start=1):
                    ws = wb[name]
                    rows = []
                    for row in ws.iter_rows():
                        cell_values = [
                            str(c.value).strip()
                            for c in row
                            if c.value is not None and str(c.value).strip()
                        ]
                        if cell_values:
                            rows.append(cell_values)
                    if rows:
                        max_col = max(len(r) for r in rows)
                        md_lines = [f"[시트: {name}]"]
                        header = rows[0] + [""] * (max_col - len(rows[0]))
                        md_lines.append("| " + " | ".join(header) + " |")
                        md_lines.append("| " + " | ".join(["---"] * max_col) + " |")
                        for row_data in rows[1:]:
                            row_data = row_data + [""] * (max_col - len(row_data))
                            md_lines.append("| " + " | ".join(row_data) + " |")
                        texts.append(("\n".join(md_lines), idx))
            except Exception:
                continue

        for text, page_num in texts:
            step = chunk_size - overlap
            start = 0
            while start < len(text):
                chunk = text[start:start + chunk_size].strip()
                if chunk:
                    documents.append({"content": chunk, "source": source, "page": page_num})
                start += step

    return documents


# 모듈 레벨 벡터스토어 캐시
_VECTORSTORE_CACHE = None


def _get_vectorstore():
    """벡터스토어 싱글톤을 반환한다."""
    global _VECTORSTORE_CACHE
    if _VECTORSTORE_CACHE is None:
        _VECTORSTORE_CACHE = _build_vectorstore()
    return _VECTORSTORE_CACHE


@tool
def search_documents(query: str):
    """사내 규정, 가이드라인, 정책 등 비정형 문서 내용을 검색합니다.

    ChromaDB 벡터 검색을 사용합니다. ChromaDB가 없으면 data/docs/ 원본
    문서를 파싱하여 자동 구축합니다.

    Args:
        query: 검색할 질문 또는 키워드 (예: "연차 사용 규정", "재택근무 조건")

    Returns:
        관련 문서 내용, 출처 파일명, 유사도 점수가 담긴 딕셔너리 목록.
    """
    logger.info("[search_documents] 검색 쿼리: %s", query)

    collection = _get_vectorstore()
    if collection is not None:
        try:
            results = collection.query(query_texts=[query], n_results=3)
            formatted = []
            if results["documents"] and results["documents"][0]:
                for doc, meta, dist in zip(
                    results["documents"][0],
                    results["metadatas"][0],
                    results["distances"][0],
                ):
                    formatted.append({
                        "content": doc,
                        "source": meta.get("source", "unknown"),
                        "score": round(1.0 - dist, 4),
                    })
            logger.info("[search_documents] 벡터 검색 결과 수: %d", len(formatted))
            return formatted
        except Exception as exc:
            logger.warning("ChromaDB 쿼리 실패: %s", exc)

    # 벡터스토어 구축 실패 시 키워드 검색 폴백
    parsed = _parse_and_chunk_docs(chunk_size=1000, overlap=0)
    if not parsed:
        return []

    query_lower = query.lower()
    scored = []
    for doc in parsed:
        score = sum(1 for w in query_lower.split() if w in doc["content"].lower())
        if score > 0:
            scored.append({**doc, "score": round(min(score * 0.1 + 0.5, 1.0), 4)})
    scored.sort(key=lambda x: x["score"], reverse=True)

    logger.info("[search_documents] 키워드 검색 결과 수: %d", len(scored[:3]))
    return scored[:3]
