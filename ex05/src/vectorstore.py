"""ex05 — ChromaDB Retriever 생성 모듈."""

import os
from pathlib import Path

from dotenv import load_dotenv
from langchain.schema import Document

load_dotenv()


def build_retriever():
    """ChromaDB Retriever를 생성한다."""
    chroma_dir = os.getenv("CHROMA_PERSIST_DIR", "./data/chroma_db")
    collection_name = os.getenv("CHROMA_COLLECTION_NAME", "metacoding_documents")
    top_k = int(os.getenv("RETRIEVER_TOP_K", "5"))
    embedding_model_name = os.getenv("EMBEDDING_MODEL", "jhgan/ko-sroberta-multitask")

    from langchain_chroma import Chroma
    from langchain_community.embeddings import HuggingFaceEmbeddings

    embeddings = HuggingFaceEmbeddings(
        model_name=embedding_model_name,
        model_kwargs={"device": "cpu"},
    )

    has_data = os.path.isdir(chroma_dir) and any(
        f.endswith(".sqlite3") for f in _list_dir_recursive(chroma_dir)
    )

    if has_data:
        print(f"[INFO] 기존 ChromaDB 로드: {chroma_dir}")
        vectorstore = Chroma(
            collection_name=collection_name,
            embedding_function=embeddings,
            persist_directory=chroma_dir,
        )
    else:
        print("[INFO] ChromaDB가 없습니다. data/docs/ 원본 문서에서 자동 구축합니다.")
        docs = _parse_and_chunk_docs()
        if not docs:
            raise RuntimeError(
                "data/docs/ 디렉토리에 문서가 없습니다. "
                "원본 문서(PDF/DOCX/XLSX)를 data/docs/에 넣고 다시 실행하십시오."
            )
        vectorstore = Chroma.from_documents(
            documents=docs,
            embedding=embeddings,
            persist_directory=chroma_dir,
            collection_name=collection_name,
        )
        print(f"[INFO] ChromaDB 자동 구축 완료: {len(docs)}건 → {chroma_dir}")

    return vectorstore.as_retriever(search_kwargs={"k": top_k})


def _list_dir_recursive(path):
    """디렉토리 내 모든 파일 이름을 재귀적으로 반환한다."""
    result = []
    for root, _, files in os.walk(path):
        result.extend(files)
    return result


def _parse_and_chunk_docs(chunk_size=500, overlap=100):
    """data/docs/의 원본 PDF/DOCX/XLSX를 파싱→청킹하여 Document 리스트로 반환한다."""
    import pypdf
    from docx import Document as DocxDocument
    import openpyxl

    docs_dir = Path(__file__).parent.parent / "data" / "docs"
    if not docs_dir.exists():
        return []

    documents = []
    for file_path in sorted(docs_dir.rglob("*")):
        suffix = file_path.suffix.lower()
        source = file_path.stem

        texts = []

        if suffix == ".pdf":
            try:
                with open(file_path, "rb") as f:
                    reader = pypdf.PdfReader(f)
                    for page_num, page in enumerate(reader.pages, start=1):
                        text = (page.extract_text() or "").strip()
                        if text:
                            md_text = f"## 페이지 {page_num}\n\n{text}"
                            texts.append((md_text, page_num))
            except Exception:
                continue

        elif suffix == ".docx":
            try:
                doc = DocxDocument(str(file_path))
                text_parts = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    style_name = para.style.name
                    if style_name == "Title":
                        text_parts.append(f"# {text}")
                    elif style_name.startswith("Heading"):
                        level_str = style_name.replace("Heading", "").strip()
                        try:
                            level = int(level_str)
                        except ValueError:
                            level = 2
                        text_parts.append(f"{'#' * level} {text}")
                    elif style_name == "List Bullet":
                        text_parts.append(f"- {text}")
                    else:
                        text_parts.append(text)
                for table in doc.tables:
                    for i, row in enumerate(table.rows):
                        row_data = [cell.text.strip() for cell in row.cells]
                        text_parts.append("| " + " | ".join(row_data) + " |")
                        if i == 0:
                            text_parts.append("| " + " | ".join(["---"] * len(row_data)) + " |")
                full_text = "\n".join(text_parts)
                if full_text:
                    texts.append((full_text, 1))
            except Exception:
                continue

        elif suffix == ".xlsx":
            try:
                wb = openpyxl.load_workbook(str(file_path), data_only=True)
                for idx, name in enumerate(wb.sheetnames, start=1):
                    ws = wb[name]
                    rows = []
                    for row in ws.iter_rows():
                        cell_values = [
                            str(c.value).strip()
                            for c in row
                            if c.value is not None and str(c.value).strip()
                        ]
                        if cell_values:
                            rows.append(cell_values)
                    if rows:
                        max_col = max(len(r) for r in rows)
                        md_lines = [f"[시트: {name}]"]
                        header = rows[0] + [""] * (max_col - len(rows[0]))
                        md_lines.append("| " + " | ".join(header) + " |")
                        md_lines.append("| " + " | ".join(["---"] * max_col) + " |")
                        for row_data in rows[1:]:
                            row_data = row_data + [""] * (max_col - len(row_data))
                            md_lines.append("| " + " | ".join(row_data) + " |")
                        texts.append(("\n".join(md_lines), idx))
            except Exception:
                continue

        for text, page_num in texts:
            step = chunk_size - overlap
            start = 0
            while start < len(text):
                chunk = text[start:start + chunk_size].strip()
                if chunk:
                    documents.append(Document(
                        page_content=chunk,
                        metadata={"source": source, "page": page_num},
                    ))
                start += step

    return documents
